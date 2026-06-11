from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_DIR = ROOT / "submission_docs"
TEMPLATE_DIR = ROOT / "作业模板" / "作业模板"

REPORTS = [
    {
        "title": "题目需求理解与分析报告",
        "source": SUBMISSION_DIR / "题目需求理解与分析报告_终稿.md",
        "template": TEMPLATE_DIR / "题目需求理解与分析报告.docx",
        "target": SUBMISSION_DIR / "题目需求理解与分析报告_终稿.docx",
    },
    {
        "title": "解题技术思路和最优方案报告",
        "source": SUBMISSION_DIR / "解题技术思路和最优方案报告_终稿.md",
        "template": TEMPLATE_DIR / "解题技术思路和最优方案报告.docx",
        "target": SUBMISSION_DIR / "解题技术思路和最优方案报告_终稿.docx",
    },
]


def clear_body(document: DocumentObject) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def style_name(document: DocumentObject, preferred: list[str], fallback: str = "Normal") -> str:
    names = {style.name for style in document.styles}
    for item in preferred:
        if item in names:
            return item
    return fallback


def add_code_paragraph(document: DocumentObject, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_markdown_line(document: DocumentObject, line: str, title_style: str, heading_style: str) -> None:
    heading = re.match(r"^(#{1,6})\s+(.+)$", line)
    if heading:
        level = len(heading.group(1))
        text = heading.group(2).strip()
        document.add_paragraph(text, style=title_style if level == 1 else heading_style)
        return

    bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
    if bullet:
        try:
            document.add_paragraph(bullet.group(1).strip(), style="List Bullet")
        except KeyError:
            document.add_paragraph("• " + bullet.group(1).strip())
        return

    numbered = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
    if numbered:
        try:
            document.add_paragraph(numbered.group(2).strip(), style="List Number")
        except KeyError:
            document.add_paragraph(line.strip())
        return

    paragraph = document.add_paragraph()
    for part in re.split(r"(`[^`]+`)", line.strip()):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def markdown_to_docx(document: DocumentObject, markdown: str, title: str) -> None:
    title_style = style_name(document, ["Title", "标题", "Heading 1", "标题 1"])
    heading_style = style_name(document, ["Heading 1", "标题 1", "Heading 2", "标题 2"])

    has_top_title = False
    in_code = False
    code_lines: list[str] = []

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                add_code_paragraph(document, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            document.add_paragraph()
            continue

        if line.startswith("# "):
            has_top_title = True
        add_markdown_line(document, line, title_style, heading_style)

    if code_lines:
        add_code_paragraph(document, "\n".join(code_lines))

    if not has_top_title:
        first = document.paragraphs[0]._p if document.paragraphs else None
        paragraph = document.add_paragraph(title, style=title_style)
        if first is not None:
            document._body._element.remove(paragraph._p)
            document._body._element.insert(0, paragraph._p)


def build_report(config: dict[str, Path | str]) -> None:
    source = Path(config["source"])
    template = Path(config["template"])
    target = Path(config["target"])
    title = str(config["title"])

    if not source.exists():
        raise FileNotFoundError(f"Missing source markdown: {source}")
    if not template.exists():
        raise FileNotFoundError(f"Missing assignment template: {template}")

    document = Document(str(template))
    clear_body(document)
    markdown_to_docx(document, source.read_text(encoding="utf-8"), title)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))
    print(f"Wrote {target.relative_to(ROOT)} using {template.relative_to(ROOT)}")


def main() -> None:
    for config in REPORTS:
        build_report(config)


if __name__ == "__main__":
    main()

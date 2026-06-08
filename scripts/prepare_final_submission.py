from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_ROOT = ROOT / "final_submission" / "Generative_Agents_课程作业提交包"

REPORTS = [
    (
        ROOT / "submission_docs" / "题目需求理解与分析报告_终稿.md",
        "题目需求理解与分析报告.docx",
    ),
    (
        ROOT / "submission_docs" / "解题技术思路和最优方案报告_终稿.md",
        "解题技术思路和最优方案报告.docx",
    ),
]

IGNORE_DIRS = {
    ".git",
    ".next",
    ".venv",
    "__pycache__",
    "node_modules",
    "final_submission",
}

IGNORE_FILES = {
    ".DS_Store",
}


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_title_page(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("课程作业提交材料\nGenerative Agents 论文复刻项目")

    doc.add_section(WD_SECTION.NEW_PAGE)


def flush_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = doc.add_paragraph()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if index < len(lines) - 1:
            run.add_break()


def add_markdown_content(doc: Document, text: str) -> None:
    image_pattern = re.compile(r"!\[.*?\]\((.*?)\)")
    numbered_pattern = re.compile(r"^\d+\.\s+")

    in_code_block = False
    code_lines: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 3)
            heading_text = stripped[level:].strip()
            doc.add_heading(heading_text, level=level)
            continue

        image_match = image_pattern.search(stripped)
        if image_match:
            image_path = Path(image_match.group(1))
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        if numbered_pattern.match(stripped):
            doc.add_paragraph(numbered_pattern.sub("", stripped), style="List Number")
            continue

        doc.add_paragraph(stripped)


def build_report(markdown_path: Path, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    title = markdown_path.stem.replace("_终稿", "")
    add_title_page(doc, title)
    add_markdown_content(doc, markdown_path.read_text(encoding="utf-8"))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def should_ignore(path: Path) -> bool:
    return path.name in IGNORE_DIRS or path.name in IGNORE_FILES


def copy_project_snapshot(target_dir: Path) -> None:
    target_dir.mkdir(parents=True, exist_ok=True)

    for item in ROOT.iterdir():
        if should_ignore(item):
            continue
        if item.name in {"submission_docs", "作业模板", "1.mp4", "frame_1.jpg", "frame_2.jpg", "frame_3.jpg", "frame_4.jpg", "frame_5.jpg"}:
            continue
        destination = target_dir / item.name
        if item.is_dir():
            shutil.copytree(
                item,
                destination,
                ignore=shutil.ignore_patterns(*IGNORE_DIRS, *IGNORE_FILES),
            )
        else:
            shutil.copy2(item, destination)


def write_submission_notes(target_dir: Path) -> None:
    note = """Generative Agents 课程作业提交说明

1. 本目录已整理为课程提交包结构。
2. `video/` 目录下暂放讲解脚本，请录制完成后替换为最终 mp4。
3. `ppt/Generative_Agents_课程答辩稿.pptx` 仍建议手动补上首页组员姓名。
4. 代码目录未包含 `.venv`、`node_modules`、`.next` 等本地依赖缓存。
5. 后端支持 OpenAI 兼容接口，也支持无 key 的稳定 fallback 演示。
"""
    (target_dir / "提交说明.txt").write_text(note, encoding="utf-8")


def build_submission_bundle() -> None:
    if SUBMISSION_ROOT.exists():
        shutil.rmtree(SUBMISSION_ROOT)

    code_dir = SUBMISSION_ROOT / "code"
    ppt_dir = SUBMISSION_ROOT / "ppt"
    report_dir = SUBMISSION_ROOT / "report"
    video_dir = SUBMISSION_ROOT / "video"
    extra_dir = SUBMISSION_ROOT / "extra"
    screenshots_dir = extra_dir / "screenshots"

    copy_project_snapshot(code_dir)

    ppt_dir.mkdir(parents=True, exist_ok=True)
    shutil.copy2(
        ROOT / "submission_docs" / "Generative_Agents_课程答辩稿.pptx",
        ppt_dir / "Generative_Agents_课程答辩稿.pptx",
    )

    report_dir.mkdir(parents=True, exist_ok=True)
    for markdown_path, docx_name in REPORTS:
        build_report(markdown_path, report_dir / docx_name)
        shutil.copy2(markdown_path, report_dir / markdown_path.name)

    video_dir.mkdir(parents=True, exist_ok=True)
    shutil.copy2(
        ROOT / "submission_docs" / "10分钟视频讲解脚本.md",
        video_dir / "10分钟视频讲解脚本.md",
    )
    (video_dir / "README_待录制.txt").write_text(
        "请将最终录制完成的 10 分钟讲解视频放在此目录，并命名为“10分钟讲解视频.mp4”。\n",
        encoding="utf-8",
    )

    screenshots_dir.mkdir(parents=True, exist_ok=True)
    for image_path in (ROOT / "submission_docs" / "screenshots").glob("*"):
        if image_path.is_file():
            shutil.copy2(image_path, screenshots_dir / image_path.name)

    extra_docs = [
        ROOT / "submission_docs" / "PPT逐页演讲稿.md",
        ROOT / "submission_docs" / "PPT逐页文案.md",
        ROOT / "submission_docs" / "答辩素材与PPT提纲.md",
        ROOT / "submission_docs" / "最终提交清单.md",
    ]
    for doc_path in extra_docs:
        shutil.copy2(doc_path, extra_dir / doc_path.name)

    write_submission_notes(SUBMISSION_ROOT)


if __name__ == "__main__":
    build_submission_bundle()
    print(f"Prepared submission bundle at: {SUBMISSION_ROOT}")
